"""Plots tab - manage saved plots."""

import asyncio
import base64
import tempfile
from datetime import datetime
from functools import partial
from pathlib import Path
from typing import Optional

import flet as ft
from jinja2 import Template

from dasmixer.api.project.project import Project
from dasmixer.gui.components.plotly_viewer import PlotlyViewer, render_png_async
from dasmixer.gui.utils import show_snack
from dasmixer.utils import logger


class PlotItemCard(ft.Container):
    """Card displaying single saved plot info."""

    def __init__(
        self,
        plot_info: dict,
        on_view,
        on_delete,
        on_select,
    ):
        super().__init__()
        self.plot_info = plot_info
        self._on_view = on_view
        self._on_delete = on_delete
        self._on_select = on_select

        # Checkbox for selection
        self.checkbox = ft.Checkbox(
            value=False,
            on_change=lambda e: on_select(plot_info["id"], e.control.value),
        )

        # Format created_at
        try:
            created_dt = datetime.fromisoformat(plot_info["created_at"])
            created_str = created_dt.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            created_str = plot_info["created_at"]

        # Extract entity_id from settings if available
        entity_id = (
            plot_info["settings"].get("entity_id", "N/A")
            if plot_info["settings"]
            else "N/A"
        )

        info_text = ft.Column(
            [
                ft.Text(
                    f"Plot #{plot_info['id']}: {plot_info['plot_type']}",
                    size=14,
                    weight=ft.FontWeight.BOLD,
                ),
                ft.Text(
                    f"Created: {created_str}",
                    size=11,
                    color=ft.Colors.GREY_600,
                ),
                ft.Text(
                    f"Entity: {entity_id}",
                    size=11,
                    color=ft.Colors.GREY_600,
                ),
            ],
            spacing=2,
            tight=True,
        )

        view_button = ft.IconButton(
            icon=ft.Icons.VISIBILITY,
            tooltip="View plot",
            on_click=lambda e: on_view(plot_info["id"]),
        )

        delete_button = ft.IconButton(
            icon=ft.Icons.DELETE,
            tooltip="Delete plot",
            icon_color=ft.Colors.RED_400,
            on_click=lambda e: on_delete(plot_info["id"]),
        )

        self.content = ft.Row(
            [
                self.checkbox,
                info_text,
                ft.Container(expand=True),
                view_button,
                delete_button,
            ],
            alignment=ft.MainAxisAlignment.START,
            vertical_alignment=ft.CrossAxisAlignment.CENTER,
        )

        self.border = ft.border.all(1, ft.Colors.GREY_300)
        self.border_radius = 5
        self.padding = 10
        self.bgcolor = ft.Colors.WHITE


class PlotsTab(ft.Container):
    """Plots tab - manage saved plots."""

    def __init__(self, project: Project):
        super().__init__()
        self.project = project
        self.expand = True
        self.padding = 10

        # State
        self.plot_items: list[dict] = []
        self.selected_ids: set[int] = set()

        # UI references
        self.plots_list: ft.ListView | None = None
        self._export_button: ft.ElevatedButton | None = None

        # Build UI
        self.content = self._build_ui()

    def _build_ui(self) -> ft.Control:
        """Build tab layout."""
        self._export_button = ft.ElevatedButton(
            content=ft.Text("Export Selected to Word"),
            icon=ft.Icons.DESCRIPTION,
            on_click=lambda e: (
                self.page.run_task(self._on_export_selected, e) if self.page else None
            ),
        )

        header = ft.Row(
            [
                ft.Text("Saved Plots", size=20, weight=ft.FontWeight.BOLD),
                ft.Container(expand=True),
                ft.ElevatedButton(
                    content=ft.Text("Select All"),
                    icon=ft.Icons.SELECT_ALL,
                    on_click=lambda e: self._on_select_all(),
                ),
                ft.ElevatedButton(
                    content=ft.Text("Deselect All"),
                    icon=ft.Icons.DESELECT,
                    on_click=lambda e: self._on_deselect_all(),
                ),
                self._export_button,
                ft.IconButton(
                    icon=ft.Icons.REFRESH,
                    tooltip="Refresh list",
                    on_click=lambda e: (
                        self.page.run_task(self.load_data) if self.page else None
                    ),
                ),
            ],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )

        self.plots_list = ft.ListView(
            spacing=10,
            padding=10,
            expand=True,
        )

        return ft.Column(
            [
                header,
                ft.Container(height=10),
                ft.Divider(),
                ft.Container(height=10),
                self.plots_list,
            ],
            expand=True,
        )

    def did_mount(self):
        """Load data when tab is mounted."""
        if self.page:
            self.page.run_task(self.load_data)

    # ------------------------------------------------------------------
    # Data loading
    # ------------------------------------------------------------------

    async def load_data(self, e=None):
        """Load list of saved plots."""
        try:
            self.plot_items = await self.project.get_saved_plots()

            self.plots_list.controls.clear()
            self.selected_ids.clear()

            if not self.plot_items:
                self.plots_list.controls.append(
                    ft.Container(
                        content=ft.Text(
                            "No saved plots yet",
                            size=14,
                            color=ft.Colors.GREY_600,
                        ),
                        alignment=ft.Alignment.CENTER,
                        padding=20,
                    )
                )
            else:
                for plot_info in self.plot_items:
                    card = PlotItemCard(
                        plot_info=plot_info,
                        on_view=lambda pid, _p=plot_info["id"]: (
                            self.page.run_task(self._view_plot, _p) if self.page else None
                        ),
                        on_delete=lambda pid, _p=plot_info["id"]: (
                            self.page.run_task(self._delete_plot, _p) if self.page else None
                        ),
                        on_select=self._on_plot_selected,
                    )
                    self.plots_list.controls.append(card)

            if self.page:
                self.page.update()

        except Exception as ex:
            logger.exception(ex)
            if self.page:
                show_snack(self.page, f"Error loading plots: {ex}", ft.Colors.RED_400)
                self.page.update()

    # ------------------------------------------------------------------
    # Selection
    # ------------------------------------------------------------------

    def _on_plot_selected(self, plot_id: int, selected: bool):
        """Handle plot selection change."""
        if selected:
            self.selected_ids.add(plot_id)
        else:
            self.selected_ids.discard(plot_id)

    def _on_select_all(self):
        """Select all plots."""
        if not self.plots_list or not self.plots_list.controls:
            return
        for card in self.plots_list.controls:
            if isinstance(card, PlotItemCard):
                card.checkbox.value = True
        # Add all plot IDs to selected_ids
        self.selected_ids = {plot_info["id"] for plot_info in self.plot_items}
        self.plots_list.update()

    def _on_deselect_all(self):
        """Deselect all plots."""
        if not self.plots_list or not self.plots_list.controls:
            return
        for card in self.plots_list.controls:
            if isinstance(card, PlotItemCard):
                card.checkbox.value = False
        self.selected_ids.clear()
        self.plots_list.update()

    # ------------------------------------------------------------------
    # View plot
    # ------------------------------------------------------------------

    async def _view_plot(self, plot_id: int):
        """View plot in dialog — loads figure, renders PNG in thread pool, then shows."""
        if not self.page:
            return

        # Show loading dialog immediately
        loading_container = ft.Container(
            content=ft.Column(
                [
                    ft.ProgressRing(width=36, height=36, stroke_width=3),
                    ft.Text("Loading plot...", size=13, color=ft.Colors.GREY_600),
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                spacing=12,
            ),
            alignment=ft.Alignment.CENTER,
            width=600,
            height=420,
        )

        close_btn = ft.TextButton(
            "Close",
            on_click=lambda e: self._close_dialog(dialog),
        )

        dialog = ft.AlertDialog(
            title=ft.Text(f"Plot #{plot_id}"),
            content=loading_container,
            actions=[close_btn],
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

        try:
            # Load figure from DB (fast — only unpickling)
            fig = await self.project.load_saved_plot(plot_id)

            # Render PNG in thread pool — Kaleido never blocks the event loop
            img_bytes = await render_png_async(fig, width=900, height=580)

            viewer = PlotlyViewer(
                figure=fig,
                width=900,
                height=580,
                title=f"Plot #{plot_id}",
                show_interactive_button=True,
                img_bytes=img_bytes,
            )

            dialog.content = ft.Container(
                content=viewer,
                width=950,
                height=640,
            )
            if self.page:
                self.page.update()

        except Exception as ex:
            logger.exception(ex)
            dialog.content = ft.Container(
                content=ft.Text(f"Error loading plot: {ex}", color=ft.Colors.RED_400),
                alignment=ft.Alignment.CENTER,
                width=500,
                height=120,
            )
            if self.page:
                self.page.update()

    # ------------------------------------------------------------------
    # Delete plot
    # ------------------------------------------------------------------

    async def _delete_plot(self, plot_id: int):
        """Delete plot with confirmation."""
        if not self.page:
            return

        async def on_confirm(e):
            dialog.open = False
            if self.page:
                self.page.update()
            await self._delete_plot_confirmed(plot_id)

        dialog = ft.AlertDialog(
            title=ft.Text("Confirm Delete"),
            content=ft.Text(f"Are you sure you want to delete plot #{plot_id}?"),
            actions=[
                ft.TextButton("Cancel", on_click=lambda e: self._close_dialog(dialog)),
                ft.ElevatedButton(
                    "Delete",
                    on_click=lambda e: (
                        self.page.run_task(on_confirm, e) if self.page else None
                    ),
                    bgcolor=ft.Colors.RED_400,
                    color=ft.Colors.WHITE,
                ),
            ],
        )

        self.page.overlay.append(dialog)
        dialog.open = True
        self.page.update()

    async def _delete_plot_confirmed(self, plot_id: int):
        """Actually delete the plot."""
        try:
            await self.project.delete_saved_plot(plot_id)
            if self.page:
                show_snack(self.page, f"Plot #{plot_id} deleted", ft.Colors.GREEN_400)
            await self.load_data()
        except Exception as ex:
            logger.exception(ex)
            if self.page:
                show_snack(self.page, f"Error deleting plot: {ex}", ft.Colors.RED_400)
                self.page.update()

    # ------------------------------------------------------------------
    # Export to Word
    # ------------------------------------------------------------------

    async def _on_export_selected(self, e):
        """Export selected plots to Word."""
        if not self.selected_ids:
            if self.page:
                show_snack(self.page, "No plots selected", ft.Colors.ORANGE_400)
                self.page.update()
            return

        try:
            # New async FilePicker API
            result = await ft.FilePicker().get_directory_path(
                dialog_title="Select folder for export"
            )
            if not result:
                return

            output_dir = Path(result)
            await self._export_plots_to_word(list(self.selected_ids), output_dir)

        except Exception as ex:
            logger.exception(ex)
            if self.page:
                show_snack(self.page, f"Error exporting: {ex}", ft.Colors.RED_400)
                self.page.update()

    async def _export_plots_to_word(self, plot_ids: list[int], output_dir: Path):
        """Export plots to Word document.

        PNG rendering is offloaded to a thread pool via render_png_async so
        the event loop stays free during Kaleido calls.
        """
        from dasmixer.gui.views.tabs.peptides.dialogs.progress_dialog import ProgressDialog
        
        dialog = None
        try:
            from docx import Document
            from docx.shared import Inches

            # Load template
            template_path = (
                Path(__file__).parent / "templates" / "plots_export.html.j2"
            )
            with open(template_path, "r", encoding="utf-8") as f:
                template = Template(f.read())

            # Show progress dialog instead of snack
            if self.page:
                dialog = ProgressDialog(self.page, "Exporting Plots to Word")
                dialog.show()
                # Убрано: show_snack(self.page, f"Exporting {len(plot_ids)} plot(s)...", ft.Colors.BLUE_400)

            # Render all figures concurrently in thread pool
            plot_info_map = {p["id"]: p for p in self.plot_items}

            async def _prepare_one(plot_id: int) -> dict | None:
                plot_info = plot_info_map.get(plot_id)
                if not plot_info:
                    return None
                try:
                    fig = await self.project.load_saved_plot(plot_id)
                    img_bytes = await render_png_async(fig, width=800, height=500)
                    png_base64 = base64.b64encode(img_bytes).decode()
                    return {
                        "id": plot_id,
                        "plot_type": plot_info["plot_type"],
                        "created_at": plot_info["created_at"],
                        "png_base64": png_base64,
                        "settings": plot_info["settings"] or {},
                        "_img_bytes": img_bytes,
                    }
                except Exception as ex:
                    logger.exception(ex)
                    return None

            # Gather all renders (asyncio.gather keeps event loop responsive
            # because each render_png_async uses run_in_executor internally)
            results = await asyncio.gather(*[_prepare_one(pid) for pid in plot_ids])
            plots_data = [r for r in results if r is not None]

            if not plots_data:
                if dialog:
                    dialog.close()
                if self.page:
                    show_snack(self.page, "No plots to export", ft.Colors.ORANGE_400)
                    self.page.update()
                return

            # Update progress after gathering
            if dialog:
                dialog.update_progress(0.5, "Building document...", f"Prepared {len(plots_data)} plots")

            # Build Word document
            doc = Document()
            doc.add_heading("DASMixer - Saved Plots", 0)
            doc.add_paragraph(
                f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            doc.add_paragraph("")

            tmp_paths: list[Path] = []
            try:
                for plot_data in plots_data:
                    doc.add_heading(
                        f"Plot #{plot_data['id']}: {plot_data['plot_type']}", 1
                    )
                    doc.add_paragraph(f"Created: {plot_data['created_at']}")

                    # Write PNG to temp file and embed
                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False
                    ) as tmp:
                        tmp.write(plot_data["_img_bytes"])
                        tmp_path = Path(tmp.name)
                    tmp_paths.append(tmp_path)

                    doc.add_picture(str(tmp_path), width=Inches(6))

                    # Settings table
                    settings = plot_data["settings"]
                    if settings:
                        doc.add_heading("Settings", 2)
                        table = doc.add_table(rows=1, cols=2)
                        table.style = "Light Grid Accent 1"
                        hdr = table.rows[0].cells
                        hdr[0].text = "Parameter"
                        hdr[1].text = "Value"
                        for key, value in settings.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(key)
                            row_cells[1].text = str(value)

                    doc.add_page_break()
            finally:
                for p in tmp_paths:
                    try:
                        p.unlink()
                    except OSError:
                        pass

            output_file = (
                output_dir
                / f"plots_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            doc.save(str(output_file))

            if dialog:
                dialog.complete(f"Saved: {output_file.name}")
                await asyncio.sleep(1)
                dialog.close()

        except Exception as ex:
            if dialog:
                dialog.close()
            logger.exception(ex)
            if self.page:
                show_snack(self.page, f"Export error: {ex}", ft.Colors.RED_400)
                self.page.update()

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _close_dialog(self, dialog: ft.AlertDialog):
        """Close dialog helper."""
        dialog.open = False
        if self.page:
            self.page.update()
